import os
import io
import math
import platform
from datetime import datetime
import streamlit as st

# Page config
st.set_page_config(page_title="AI-Transitie Toolkit", layout="wide")

# --- (Optioneel) OpenAI client ---
try:
    from openai import OpenAI
    _OPENAI_AVAILABLE = True
except Exception:
    _OPENAI_AVAILABLE = False

# --- (Optioneel) DOCX export ---
try:
    from docx import Document
    from docx.shared import Pt
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False

# ========= Helpers =========
def approx_tokens(txt: str) -> int:
    if not txt:
        return 0
    return max(1, math.ceil(len(txt) / 4))

def make_docx_bytes(title: str, body_md: str) -> bytes:
    """Zet eenvoudige tekst/markdown om naar een .docx-bestand (basisopmaak)."""
    doc = Document()
    _ = doc.add_heading(title, level=1)
    for line in body_md.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def get_openai_api_key() -> str | None:
    key = None
    try:
        if 'OPENAI_API_KEY' in st.secrets:
            key = st.secrets['OPENAI_API_KEY']
    except Exception:
        pass
    if not key:
        key = os.environ.get("OPENAI_API_KEY")
    return key

def status_badge(ok: bool) -> str:
    return "✅" if ok else "⚠️"

# ===== Bouw prompt en (optioneel) roep OpenAI API aan =====
def genereer_ai_advies(software, knelpunten, weging, toelichting, voorkeur, model: str, temperature: float = 0.4):
    punten = [f"- {k} (weging {weging.get(k, 3)})" for k in knelpunten]
    beschrijving = "\n".join(punten) if punten else "- (geen geselecteerde knelpunten)"

    system_msg = (
        "Je bent een nuchtere AI-implementatiecoach voor administratie- en belastingadvieskantoren. "
        "Geef concrete, korte adviezen in duidelijk Nederlands; vermijd jargon; toon stappen en quick wins; "
        "wees specifiek per software en knelpunt. Voeg indien relevant risico's en mitigaties toe."
    )
    user_msg = f"""
Kantoorsoftware: {software}
Knelpunten en weging:\n{beschrijving}
Toelichting: {toelichting or 'n.v.t.'}
Voorkeursaanpak: {voorkeur}

Maak een voorstel met:
1) Diagnose: analyse van de huidige situatie
2) Top 3 quick wins (met inschatting tijd/impact)
3) Aanpak per fase (diagnose, PoC, opschaling) met concrete acties
4) Benodigde tooling (heel concreet)
5) Indicatieve tijdslijn (weken) en betrokken rollen
6) Risico's + mitigatie
7) Heldere vervolgstap en call-to-action voor de klant
"""

    api_key = get_openai_api_key()

    if not _OPENAI_AVAILABLE or not api_key:
        advies_stub = (
            "⚠️ Geen OpenAI-API geconfigureerd. Hieronder de samengestelde opdracht die naar het model wordt gestuurd.\n\n"
            f"**Systeemrol**:\n{system_msg}\n\n**Gebruikersinvoer**:\n{user_msg.strip()}\n\n"
            "👉 Stel OPENAI_API_KEY in via .streamlit/secrets.toml of als omgevingsvariabele."
        )
        return advies_stub

    try:
        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg},
            ],
            temperature=temperature,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return f"❌ Fout bij genereren van advies ({type(e).__name__}): {e}"

# ================= Sidebar =================
st.sidebar.title("AI-Transitie Toolkit")
menu = st.sidebar.radio("Navigatie", [
    "Dashboard",
    "Keuzeadvies",
    "Transitieplanner",
    "Proof-of-Concepts",
    "Feedback",
    "Handleiding",
    "Chatbot",  # nieuw
])

with st.sidebar.expander("Model & instellingen", expanded=False):
    st.caption("Deze instellingen gelden voor AI-advies (indien API-key ingesteld).")
    model = st.selectbox("Model", ["gpt-4o-mini", "gpt-4o"], index=0)
    temperature = st.slider("Creativiteit (temperature)", 0.0, 1.0, 0.4, 0.1)
    st.session_state['model'] = model
    st.session_state['temperature'] = temperature

with st.sidebar.expander("AI-instellingen", expanded=True):
    ai_enabled = st.checkbox("AI inschakelen (OpenAI)", value=True, help="Schakel uit om gratis te testen zonder API-verbruik. Toont dan de samengestelde prompt i.p.v. AI-tekst.")
    st.session_state['ai_enabled'] = ai_enabled

with st.sidebar.expander("Systeemcheck", expanded=True):
    st.write("OpenAI lib:", status_badge(_OPENAI_AVAILABLE))
    st.write("python-docx:", status_badge(_DOCX_AVAILABLE))
    st.write("API-key:", status_badge(bool(get_openai_api_key())))
    st.write("AI ingeschakeld:", "✅" if st.session_state.get('ai_enabled', True) else "⛔")
    st.caption(f"Python {platform.python_version()} | Platform: {platform.system()}")

# Resetknop
if st.sidebar.button("🔄 Reset invoer & voorstel"):
    for k in ["software", "knelpunten", "toelichting", "weging", "voorkeur", "advies_tekst", "advies_timestamp"]:
        if k in st.session_state:
            del st.session_state[k]
    st.sidebar.success("Invoer gewist. Ga naar Dashboard om opnieuw te starten.")

# ================= Pagina's =================
# Chat state initialiseren
if "chat_history" not in st.session_state:
    st.session_state["chat_history"] = []  # lijst met {"role": "user"|"assistant", "content": str}

if menu == "Dashboard":
    st.title("AI-Transitie Toolkit voor Administratiekantoren")
    st.write("Doorloop een traject om AI-oplossingen binnen uw kantoor te integreren.")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Huidige Software en Knelpunten")
        software = st.selectbox("Selecteer software", ["Exact", "Nextens", "SnelStart", "AFAS", "Twinfield", "Multivers", "Minox", "Anders"])
        if software == "Anders":
            extra = st.text_input("Andere software (specificeer)")
            if extra:
                software = extra

        alle_knelpunten = [
            "Handmatig werk",
            "Trage communicatie",
            "Foutgevoeligheid",
            "Gebrekkig overzicht",
            "Duplicatie van gegevens",
            "Moeizame rapportages",
            "Veel tijd kwijt aan controles",
            "Onvoldoende koppelingen met andere software",
            "Compliance risico's",
        ]
        knelpunten = st.multiselect("Selecteer knelpunten", alle_knelpunten)
        toelichting = st.text_area("Toelichting of andere knelpunten")

        st.subheader("Weging van de problemen (1 = klein, 5 = groot)")
        probleemweging = {k: 3 for k in knelpunten}
        for k in knelpunten:
            probleemweging[k] = st.slider(f"Weging voor: {k}", 1, 5, probleemweging[k])

        valid = True
        if len(knelpunten) == 0:
            st.warning("Selecteer minstens één knelpunt.")
            valid = False
        for k in knelpunten:
            if probleemweging.get(k) is None:
                st.warning(f"Weging ontbreekt voor: {k}")
                valid = False

        if st.button("Advies voorbereiden"):
            if not valid:
                st.error("Corrigeer de invoer voordat je doorgaat.")
            else:
                st.session_state['software'] = software
                st.session_state['knelpunten'] = knelpunten
                st.session_state['toelichting'] = toelichting
                st.session_state['weging'] = probleemweging
                st.success("Invoer opgeslagen. Ga naar 'Keuzeadvies' om het voorstel te genereren.")

    with col2:
        st.subheader("AI-Oplossingen en Voorbeelden")
        st.markdown("""
            - **API-koppelingen** voor geautomatiseerde data-uitwisseling
            - **ChatGPT** voor diagnose en advies op maat
            - **OCR & RPA** voor documentverwerking en automatisering
            - **Zapier of Make** voor workflow-automatisering
            """)
        st.subheader("Invoerstatus")
        st.write("Software:", st.session_state.get('software', '—'))
        st.write("Knelpunten:", ", ".join(st.session_state.get('knelpunten', [])) or "—")
        st.write("Toelichting:", st.session_state.get('toelichting', '—') or "—")
        w = st.session_state.get('weging', {})
        if w:
            st.write("Weging:")
            for k, v in w.items():
                st.write(f"- {k}: {v}")
        else:
            st.write("Weging: —")

elif menu == "Keuzeadvies":
    st.title("Keuzeadvies voor jouw AI-transitie")
    st.write("Kies de aanpak die het beste past bij jouw situatie.")

    voorkeur = st.radio("Hoe wil je jouw AI-transitie aanpakken?", ["Zelf Doen", "Laten Doen", "Hybride"], index=2)
    st.session_state['voorkeur'] = voorkeur

    if voorkeur == "Zelf Doen":
        st.info("Maximale controle. Wij leveren tools, templates en begeleiding voor zelfstandig werken.")
    elif voorkeur == "Laten Doen":
        st.info("Volledige ontzorging. Een expert implementeert de AI-oplossingen voor jou.")
    else:
        st.info("Combinatie: zelf grip + hulp bij technische of complexe onderdelen.")

    missing = []
    if not st.session_state.get('software'):
        missing.append("software")
    if not st.session_state.get('knelpunten'):
        missing.append("knelpunten")
    if missing:
        st.warning("Ontbrekende invoer: " + ", ".join(missing) + ". Ga terug naar het Dashboard.")

    if st.button("Genereer Diagnose & Transitievoorstel (AI)"):
        if missing:
            st.error("Kan geen advies genereren; vul eerst de ontbrekende velden in op het Dashboard.")
        else:
            use_ai = st.session_state.get('ai_enabled', True) and bool(get_openai_api_key())
            if use_ai:
                advies_tekst = genereer_ai_advies(
                    software=st.session_state.get('software'),
                    knelpunten=st.session_state.get('knelpunten', []),
                    weging=st.session_state.get('weging', {}),
                    toelichting=st.session_state.get('toelichting', ''),
                    voorkeur=st.session_state.get('voorkeur', 'Hybride'),
                    model=st.session_state.get('model', 'gpt-4o-mini'),
                    temperature=st.session_state.get('temperature', 0.4),
                )
            else:
                _kn = st.session_state.get('knelpunten', [])
                _wg = st.session_state.get('weging', {})
                beschrijving = "\n".join([f"- {k} (weging {_wg.get(k, '—')})" for k in _kn]) if _kn else "- (geen geselecteerde knelpunten)"

                advies_tekst = (
                    "🔒 AI staat uit of er is geen API-key.\n\n"
                    f"**Software:** {st.session_state.get('software','—')}\n"
                    f"**Knelpunten en weging:**\n{beschrijving}\n"
                    f"**Toelichting:** {st.session_state.get('toelichting','—')}\n"
                    f"**Voorkeursaanpak:** {st.session_state.get('voorkeur','Hybride')}\n\n"
                    "Dit is de samengestelde opdracht (prompt) die normaal naar het model zou gaan."
                )

            st.session_state['advies_tekst'] = advies_tekst
            st.session_state['advies_timestamp'] = datetime.now().strftime("%Y-%m-%d %H:%M")
            st.subheader("Diagnose & Voorstel voor jouw AI-transitie")
            st.markdown(advies_tekst)

    if st.session_state.get('advies_tekst'):
        st.divider()
        st.subheader("Download jouw voorstel")
        titel = f"AI-Transitie Voorstel - {st.session_state.get('software','n.v.t.')} - {st.session_state.get('advies_timestamp','')}"

        md_bytes = st.session_state['advies_tekst'].encode('utf-8')
        st.download_button(
            label="Download als Markdown (.md)",
            data=md_bytes,
            file_name=f"{titel}.md",
            mime="text/markdown",
        )

        if _DOCX_AVAILABLE:
            docx_bytes = make_docx_bytes(titel, st.session_state['advies_tekst'])
            st.download_button(
                label="Download als Word (.docx)",
                data=docx_bytes,
                file_name=f"{titel}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            st.info("Voor .docx-export installeer: `pip install python-docx`. Daarna de app herstarten.")

    with st.expander("Kosteninschatting (ruw)"):
        prompt_preview = (
            str(st.session_state.get('software', ''))
            + str(st.session_state.get('knelpunten', ''))
            + str(st.session_state.get('weging', ''))
            + str(st.session_state.get('toelichting', ''))
            + str(st.session_state.get('voorkeur', ''))
        )
        in_tok = approx_tokens(prompt_preview)
        out_tok = st.slider("Verwachte output-tokens", 400, 2000, 900, 50)
        st.write(f"Geschatte input tokens: ~{in_tok}")
        st.write(f"Geschatte output tokens: ~{out_tok}")
        st.caption("Let op: dit is een ruwe schatting. Werkelijke aantallen kunnen afwijken.")

elif menu == "Chatbot":
    st.title("Chatbot – Project Kantoor in AI-Transitie")

    # Toon historie
    for msg in st.session_state["chat_history"]:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])    

    # Invoer
    prompt = st.chat_input("Stel je vraag over software, routines of AI‑kansen…")
    if prompt:
        # Voeg userbericht toe en toon het
        st.session_state["chat_history"].append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        use_ai = st.session_state.get("ai_enabled", True) and bool(get_openai_api_key()) and _OPENAI_AVAILABLE
        if use_ai:
            system_msg = (
                "Je bent een nuchtere, praktische AI-coach voor administratie- en belastingadvieskantoren. "
                "Geef korte, concrete antwoorden in helder Nederlands; vermijd jargon; noem waar zinvol stappen/voorbeelden. "
                "Houd rekening met software als Exact, Nextens, SnelStart, AFAS, Twinfield, Multivers, Minox."
            )
            try:
                from openai import OpenAI
                client = OpenAI(api_key=get_openai_api_key())
                # Beperk context voor kostenbeheersing
                ctx = st.session_state["chat_history"][-8:]
                resp = client.chat.completions.create(
                    model=st.session_state.get("model", "gpt-4o-mini"),
                    temperature=st.session_state.get("temperature", 0.4),
                    messages=[{"role": "system", "content": system_msg}] + ctx,
                    max_tokens=600,
                )
                answer = resp.choices[0].message.content.strip()
            except Exception as e:
                answer = f"❌ Fout bij genereren van antwoord: {e}"
        else:
            # Gratis/stub modus: hulptekst + echo
            answer = (
                "🔒 AI staat uit of er is geen API‑key.

"
                "Ik kan je vraag wel structureren. Probeer bijvoorbeeld:
"
                "- Wat is de snelste manier om [knelpunt] te verminderen in [software]?
"
                "- Geef een stappenplan voor [proces] met een quick win.

"
                f"Je vroeg: **{prompt}**"
            )

        with st.chat_message("assistant"):
            st.markdown(answer)
        st.session_state["chat_history"].append({"role": "assistant", "content": answer})

    st.caption("Tip: gebruik de resetknop in de sidebar om het gesprek te wissen.")

elif menu == "Transitieplanner":
    st.title("Transitieplanner")
    st.info("Deze sectie wordt in de volgende versie uitgewerkt.")

elif menu == "Proof-of-Concepts":
    st.title("Proof-of-Concepts")
    st.info("Hier komen uitgewerkte voorbeelden met instructies.")

elif menu == "Feedback":
    st.title("Feedback")
    feedback = st.text_area("Wat vond je van deze toolkit?")
    if st.button("Verzenden"):
        st.success("Bedankt voor je feedback!")

elif menu == "Handleiding":
    st.title("Handleiding voor installatie en gebruik")
    st.markdown("""
        ### 1. Waar plaats je de bestanden?
        Plaats de bestanden **niet in C:\\Python313** maar in een eigen map, bijvoorbeeld:
        ```
        C:\\Users\\<jouwnaam>\\AI-Toolkit\\
        ```

        ### 2. Virtuele omgeving maken (aanbevolen)
        ```powershell
        cd C:\\Users\\<jouwnaam>\\AI-Toolkit
        python -m venv venv
        venv\\Scripts\\activate
        ```

        ### 3. Dependencies installeren
        ```powershell
        pip install streamlit openai python-docx
        ```

        ### 4. API-sleutel instellen
        **Optie A (snel, alleen huidig venster):**
        ```powershell
        $env:OPENAI_API_KEY = "sk-...jouw_key..."
        ```
        **Optie B (permanent voor jouw account):**
        ```powershell
